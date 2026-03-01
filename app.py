import streamlit as st
import io
import os

from transpose_harmony import transpose_harmony

#streamlit run app.py

def read_file(uploaded_file) -> str:
    ext = os.path.splitext(uploaded_file.name)[1].lower()

    if ext == ".txt":
        return uploaded_file.read().decode("utf-8")

    elif ext == ".pdf":
        import fitz  # pip install pymupdf
        data = uploaded_file.read()
        doc = fitz.open(stream=data, filetype="pdf")
        return "\n".join(page.get_text() for page in doc)

    elif ext == ".docx":
        from docx import Document  # pip install python-docx
        doc = Document(uploaded_file)
        return "\n".join(para.text for para in doc.paragraphs)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


def write_file(text: str, original_filename: str):
    ext = os.path.splitext(original_filename)[1].lower()

    if ext == ".txt":
        return text.encode("utf-8"), "text/plain"

    elif ext == ".pdf":
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), text, fontsize=11)
        return doc.tobytes(), "application/pdf"

    elif ext == ".docx":
        from docx import Document
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ── UI ──────────────────────────────────────────

st.set_page_config(page_title="Append Number to File", page_icon="🔢")
st.title("🔢 Append Number to File")
st.write("Upload a file and enter requested tranposition:")

uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx"])
shift = st.number_input("Transpose (semitones): ", step=1)

if uploaded_file and st.button("Process"):
    try:
        harmony_string = read_file(uploaded_file)
        result_text = transpose_harmony(harmony_string, shift)

        with st.expander("Result preview"):
            st.text(result_text)

        output_bytes, mime = write_file(result_text, uploaded_file.name)
        input_name = uploaded_file.name
        ext = input_name.split('.')[-1].lower()
        output_name = '.'.join(input_name.split('.')[:-1])
        output_name +='-shifted_by' + str(shift) + '.'+ ext
        st.download_button(
            label="⬇️ Download file",
            data=output_bytes,
            file_name=output_name,
            mime=mime
        )

    except Exception as e:
        st.error(f"Error: {e}")
