import numpy as np

white_notes = [57,59,60,62,64,65,67]
#chord_names =['Am','E','Dm', 'F', 'C', 'G']
#extended_chord_names =['Am','E','Dm', 'F', 'C','G','A','Em', 'D','G#o', 'Bo', 'Bb', 'Fm', 'B']

bases = [['Ab','G#'], ['A'], ['Bb','A#'], ['B','Cb'], ['C','B#'], ['Db','C#'], ['D'], ['Eb','D#'], ['E'], ['F'], ['F#','Gb'], ['G']]
#kind = ['','m','7','m7', 'dim', '4', 'sus', 'o7', '7j', 'm7j', 'O']


def transpose_harmony(harmony_string, shift):
    output_string = ''
    text_index = 0
    while text_index!=len(harmony_string):
        text_base =  [harmony_string[text_index:text_index+2], harmony_string[text_index]]
        """ handle spaces case, optional ..."""
        found = False
        for b in text_base:
            for b_index in range(0,12):
                b_group  = bases[b_index]
                if b in b_group:
                    new_b_index = (b_index + shift) % 12
                    new_b = bases[new_b_index][0]
                    output_string += new_b
                    text_index += len(b)
                    found = True
                    break
            if (found): 
                break
        if not found:
            output_string += harmony_string[text_index]
            text_index += 1
    return (output_string)

def transpose_harmony_in_txt(input_file, output_file, shift):
    input_string = open(input_file, "r", encoding="utf-8").read()
    output_string = transpose_harmony(harmony_string, shift)
    f = open(output_file, "w", encoding="utf-8")
    f.write(output_string)
    f.close()
    
from docx import Document

def transpose_harmony_in_docx(input_file, output_file, shift):

    doc = Document(input_file)
    
    for para in doc.paragraphs:
        for run in para.runs:
            old_harmony_string = run.text
            new_harmony_string = transpose_harmony(old_harmony_string, shift)
            run.text = new_harmony_string
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        old_harmony_string = run.text
                        new_harmony_string = transpose_harmony(old_harmony_string, shift)
                        run.text = new_harmony_string    
    doc.save(output_file)
    
import fitz  # pymupdf

def transpose_harmony_in_pdf(input_file, output_file, shift):
    doc = fitz.open(input_file)
    for page in doc:
        blocks = page.get_text("blocks")
        
        for block in blocks:
            if block[6] != 0:
                continue            
            old_harmony_string = block[4].strip()
            new_harmony_string = transpose_harmony(old_harmony_string, shift)
            if new_harmony_string != old_harmony_string:
                rect = fitz.Rect(block[:4])
                page.add_redact_annot(rect)
                page.apply_redactions()
                page.insert_text(
                    (rect.x0, rect.y1),
                    new_harmony_string,
                    fontsize=11
                )
    
    doc.save(output_file)

import os


def file_transpose_harmony(input_file, shift):
    
    ext = input_file.split('.')[-1].lower()
    
    output_file = '.'.join(input_file.split('.')[:-1])
    output_file +='-shifted_by' + str(shift) + '.'+ ext
    print(input_file, output_file)
    if ext == "txt":
        transpose_harmony_in_txt(input_file, output_file, shift)
    elif ext == "pdf":
        transpose_harmony_in_pdf(input_file, output_file, shift)
    elif ext == "docx":
        transpose_harmony_in_docx(input_file, output_file, shift)
    else:
        """ don't support file """
